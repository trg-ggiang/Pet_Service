from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = r"D:\Pet_Sevice\output\documents\Chuong_7_Nguyen_tac_thiet_ke_Pet_Service.docx"

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E2F3"
TEXT = "222222"
MUTED = "666666"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def set_font(run, size=13, bold=False, italic=False, color=TEXT, name="Times New Roman"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.3, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_body(doc, text, bold_lead=None, after=7, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=after, line=1.3, align=align)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_font(r2)
    else:
        set_font(p.add_run(text))
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    set_paragraph_spacing(p, after=4, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.left_indent = Cm(0.65 + level * 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    for run in p.runs:
        set_font(run, size=12.5)
    if not p.runs:
        set_font(p.add_run(text), size=12.5)
    else:
        p.runs[0].text = text
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(p, after=4, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    for run in p.runs:
        set_font(run, size=12.5)
    if not p.runs:
        set_font(p.add_run(text), size=12.5)
    else:
        p.runs[0].text = text
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    sizes = {1: 16, 2: 14, 3: 13}
    before = {1: 16, 2: 12, 3: 9}
    after = {1: 8, 2: 6, 3: 4}
    set_paragraph_spacing(p, before=before[level], after=after[level], line=1.1)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.outline_level = level - 1
    set_font(p.add_run(text), size=sizes[level], bold=True, color=NAVY if level < 3 else TEXT)
    return p


def add_label(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=5, after=3, line=1.15)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), size=13, bold=True, color=BLUE)
    return p


def add_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    table.allow_autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Cm(widths_cm[i])
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=0, line=1.05)
        set_font(p.add_run(text), size=10.5, bold=True, color=WHITE)
    for row_idx, row_values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].width = Cm(widths_cm[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_idx % 2 == 1:
                set_cell_shading(cells[i], LIGHT_GRAY)
            p = cells[i].paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.1, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_font(p.add_run(str(value)), size=10.5)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=2, line=1)
    return table


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=3, line=1.15)
    set_font(p.add_run(title + ": "), size=11.5, bold=True, color=NAVY)
    set_font(p.add_run(text), size=11.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.2)
section.header_distance = Cm(1.1)
section.footer_distance = Cm(1.1)

# Academic-report override of the narrative_proposal preset: A4 + Times New Roman.
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
normal.font.size = Pt(13)
normal.font.color.rgb = RGBColor.from_string(TEXT)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.3

for style_name in ("List Bullet", "List Bullet 2", "List Number"):
    style = styles[style_name]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style.font.size = Pt(12.5)

# Running header/footer.
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_paragraph_spacing(hp, after=0, line=1)
set_font(hp.add_run("PET SERVICE  |  NGUYÊN TẮC THIẾT KẾ"), size=9, bold=True, color=MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(fp, after=0, line=1)
set_font(fp.add_run("Báo cáo phân tích dựa trên mã nguồn hiện tại"), size=9, italic=True, color=MUTED)

# Editorial opening block.
p = doc.add_paragraph()
set_paragraph_spacing(p, before=35, after=14, line=1, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(p.add_run("BÁO CÁO THIẾT KẾ PHẦN MỀM"), size=11, bold=True, color=BLUE)
p = doc.add_paragraph()
set_paragraph_spacing(p, after=8, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(p.add_run("CHƯƠNG 7. NGUYÊN TẮC THIẾT KẾ"), size=25, bold=True, color=NAVY)
p = doc.add_paragraph()
set_paragraph_spacing(p, after=22, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(p.add_run("Đánh giá Coupling, Cohesion và các nguyên lý SOLID trong hệ thống Pet Service"), size=14, italic=True, color=MUTED)

add_callout(
    doc,
    "Phạm vi đánh giá",
    "Frontend React/Vite/TypeScript và backend Express theo cấu trúc domain module hiện tại, bao gồm cả đợt tái cấu trúc đang có trong thư mục làm việc. Nội dung được đối chiếu với nghiệp vụ của báo cáo ITSSfinal2: khách hàng, nhân viên, bác sĩ, quản trị viên; đặt lịch, khám bệnh, grooming, boarding, hóa đơn và thông báo.",
)

add_heading(doc, "7.1. Áp dụng Design Concepts: Coupling và Cohesion", 1)
add_heading(doc, "7.1.1. Giới thiệu chung", 2)
add_body(doc, "Cohesion (tính kết dính) phản ánh mức độ các thành phần bên trong một module cùng phục vụ một mục tiêu rõ ràng. Module có cohesion cao thường dễ đọc, dễ kiểm thử và có ít lý do thay đổi. Ví dụ, một service chỉ xử lý hồ sơ thú cưng sẽ có cohesion tốt hơn một file đồng thời xử lý thú cưng, hóa đơn, gửi email và định dạng giao diện.")
add_body(doc, "Coupling (tính liên kết) phản ánh mức độ phụ thuộc giữa các module. Coupling thấp giúp thay đổi một thành phần mà ít làm lan truyền sửa đổi sang thành phần khác. Trong ứng dụng web, service API giúp UI không phụ thuộc trực tiếp vào cơ sở dữ liệu; middleware giúp route không lặp lại xác thực; module registry giúp app không phải biết chi tiết của từng domain.")
add_body(doc, "Mục tiêu thiết kế không phải là loại bỏ mọi phụ thuộc. Một hệ thống hoạt động luôn cần các module phối hợp với nhau. Mục tiêu đúng là duy trì phụ thuộc có chủ đích, thông qua contract ổn định, đồng thời giữ mỗi module tập trung vào một nhóm trách nhiệm có quan hệ chặt chẽ.")

add_heading(doc, "7.1.2. Đánh giá phía Frontend", 2)
add_heading(doc, "7.1.2.1. Cohesion của Frontend", 3)
add_label(doc, "Ưu điểm")
for item in [
    "Cấu trúc feature theo vai trò rõ ràng: features/admin, features/auth, features/customer, features/doctor và features/staff. Cách chia này phù hợp trực tiếp với bốn tác nhân chính trong báo cáo ITSSfinal2.",
    "Thành phần trình bày được chia theo domain trong components/admin, components/customer, components/doctor và components/staff; bộ thành phần dùng chung nằm trong components/ui. Việc tách này giúp các component nhỏ tập trung vào một khu vực giao diện.",
    "Các service khách hàng được tách theo nghiệp vụ như customerAppointmentsApi, customerBoardingApi, customerMedicationsApi, customerNotificationsApi và customerPetsApi. TypeScript type tương ứng cũng được đặt trong types/customer, giúp logic truy cập API và contract dữ liệu có phạm vi rõ.",
    "Các trang feature thường đóng vai trò container, còn phần hiển thị chi tiết được đưa xuống component. Ví dụ các trang admin gọi service rồi chuyển dữ liệu cho AdminDashboardView, AdminReportsView hoặc AdminSettingsView.",
]: add_bullet(doc, item)

add_label(doc, "Hạn chế")
for item in [
    "Một số component vẫn chứa quá nhiều trạng thái, modal và luồng nghiệp vụ. Các điểm nóng tiêu biểu là StaffBoardingTab.tsx (khoảng 1.198 dòng), CustomerAppointmentModals.tsx (khoảng 977 dòng), AdminStaffView.tsx (khoảng 954 dòng) và CustomerPortal.tsx (khoảng 943 dòng). Khi thêm chức năng, các file này có nguy cơ trở thành God Component.",
    "App.tsx đồng thời quản lý vòng đời session, chuyển màn hình xác thực, phân vai trò và toàn bộ điều hướng của admin. Quy mô hiện tại vẫn kiểm soát được nhưng trách nhiệm sẽ tăng nhanh khi có deep-link, route guard hoặc nhiều layout.",
    "Một số component chuyên khoa của doctor/staff tự gọi fetch và tự xử lý lỗi thay vì đi qua service thống nhất. Điều này làm logic giao tiếp HTTP chen vào component trình bày.",
    "Các kiểu dữ liệu và phương thức của doctorAppointmentsService, staffAppointmentsService khá lớn; trong cùng một file có cả contract, mapping mặc định và nhiều nhóm use case khác nhau.",
]: add_bullet(doc, item)

add_label(doc, "Đề xuất nâng cao Cohesion")
for item in [
    "Tách component lớn theo mô hình container + view + modal; ví dụ StaffBoardingTab thành BoardingList, BoardingRoomManager, BoardingCheckoutDialog và useBoardingOperations.",
    "Tách session/role routing khỏi App.tsx bằng AuthProvider, RoleGuard và các layout riêng; sử dụng router config khi cần URL thật cho từng màn hình.",
    "Di chuyển mọi thao tác HTTP khỏi component sang service hoặc hook chuyên dụng, đồng thời tái sử dụng requestJson cho xác thực, parse lỗi và cấu hình header.",
    "Tách contract TypeScript khỏi implementation service ở các file lớn, ví dụ doctorAppointments.types.ts và doctorAppointments.api.ts.",
]: add_bullet(doc, item)

add_heading(doc, "7.1.2.2. Coupling của Frontend", 3)
add_label(doc, "Ưu điểm")
for item in [
    "UI không truy cập Supabase trực tiếp. Dữ liệu đi qua service frontend rồi đến Express API, nhờ đó giao diện không phụ thuộc vào câu truy vấn hoặc cấu trúc client database.",
    "apiUrl.ts gom cách xác định backend origin; authSession.ts gom session và auth header; requestJson.ts chuẩn hóa request JSON và xử lý lỗi. Đây là các điểm nối dùng chung giúp giảm coupling với fetch nguyên thủy.",
    "Các file service tương thích ở services/doctorAppointments.ts và services/staffAppointments.ts chỉ re-export implementation trong feature, giúp giữ ổn định import cũ khi tái cấu trúc.",
    "TypeScript DTO tạo contract tường minh giữa page, component và API, đặc biệt trong nhóm customer appointments, medications, notifications và pets.",
]: add_bullet(doc, item)

add_label(doc, "Hạn chế")
for item in [
    "Cách gọi HTTP chưa thống nhất: customer dùng requestJson, trong khi doctor và staff có các bản fetchWithAuth riêng; một số component tiếp tục gọi fetch trực tiếp. Nếu format lỗi hoặc cơ chế token thay đổi, nhiều nơi phải sửa.",
    "Frontend phụ thuộc trực tiếp vào URL endpoint dạng chuỗi và shape response cụ thể. Chưa có lớp client được sinh từ schema hoặc contract dùng chung giữa frontend và backend.",
    "Backend trả về một số dữ liệu thiên về hiển thị như class Tailwind cls/dot và màu sắc status. Cách này giảm mapping tại frontend nhưng tạo coupling ngược: backend biết chi tiết theme của UI.",
    "App.tsx dùng điều kiện role và state nội bộ để chuyển portal. Khi thêm vai trò hoặc route, file trung tâm phải thay đổi, thể hiện coupling tương đối cao giữa app shell và danh sách portal.",
]: add_bullet(doc, item)

add_label(doc, "Đề xuất giảm Coupling")
for item in [
    "Chuẩn hóa một ApiClient duy nhất có auth, timeout, parse JSON, refresh/invalid-session và ApiError.",
    "Dùng contract trung lập cho trạng thái, ví dụ statusCode/statusLabel; class CSS và màu sắc nên được map trong design system của frontend.",
    "Tạo route/portal registry theo role thay cho chuỗi if liên tiếp; mỗi portal đăng ký layout và entry component.",
    "Nếu dự án tiếp tục mở rộng, dùng OpenAPI hoặc schema validation dùng chung để giảm sai lệch DTO giữa hai phía.",
]: add_bullet(doc, item)

add_heading(doc, "7.1.3. Đánh giá phía Backend", 2)
add_heading(doc, "7.1.3.1. Cohesion của Backend", 3)
add_label(doc, "Ưu điểm")
for item in [
    "Đợt tái cấu trúc hiện tại đã tách app.js khỏi server.js. app.js chỉ tạo Express app và mount module; server.js chỉ mở cổng và khởi động reminder scheduler. Đây là cải thiện rõ rệt về cohesion và khả năng test/deploy.",
    "Mã nguồn runtime được nhóm theo domain trong modules/auth, customers, pets, appointments, doctors, medical, staff, boarding, billing, notifications và users. Cách tổ chức này bám sát các nghiệp vụ trong báo cáo thay vì chia thuần túy theo loại kỹ thuật.",
    "Các yếu tố cắt ngang đã có vị trí riêng: config cho môi trường/CORS, middlewares cho xác thực và phân quyền, lib cho Prisma/Supabase/JWT, utils cho helper HTTP.",
    "modules/index.js đóng vai trò composition root, khai báo tập trung prefix và router. Việc bổ sung một module mới không làm app.js biết chi tiết handler.",
    "Các route auth và phần lớn route customer/staff gọi service chuyên biệt, giúp handler tập trung vào request, response và status code.",
]: add_bullet(doc, item)

add_label(doc, "Hạn chế")
for item in [
    "medical.routes.js khoảng 1.983 dòng và còn chứa nhiều helper, truy vấn Supabase, mapping dữ liệu và nghiệp vụ khám. File này vừa là router vừa gần như controller/service, nên cohesion thấp dù đã nằm đúng domain.",
    "Một số service domain rất lớn: admin.service.js khoảng 1.747 dòng, staff.service.js khoảng 1.386 dòng, customer-appointment.service.js khoảng 1.054 dòng, boarding.service.js khoảng 1.028 dòng và doctor.service.js khoảng 981 dòng. Mỗi file gom nhiều use case có nhịp thay đổi khác nhau.",
    "Ranh giới một số module mới chỉ mang tính facade. reports/index.js re-export hàm từ users/admin.service.js; grooming/index.js re-export hàm từ staff/staff.service.js. Domain đã có tên riêng nhưng chưa sở hữu logic của mình.",
    "Route doctor, medical, service và staff vẫn có nơi import Supabase trực tiếp. Điều này làm quy tắc 'route mỏng, business logic ở service' chưa được áp dụng đồng đều.",
    "Logic sửa mojibake, label và view model xuất hiện trong medical.service.js; đây là nhiều loại trách nhiệm nằm chung với truy vấn lịch khám.",
]: add_bullet(doc, item)

add_label(doc, "Đề xuất nâng cao Cohesion")
for item in [
    "Tách medical.routes.js thành router mỏng và các controller/use-case service: exam-session, prescription, specialist-order, doctor-notification.",
    "Chia admin.service.js theo nhóm dashboard, user management, service catalog, staff scheduling, reports và settings; để reports/grooming/boarding thực sự sở hữu nghiệp vụ tương ứng.",
    "Tách mapper/formatter/DTO khỏi service truy vấn; ví dụ medical.mapper.js và doctor-portal.presenter.js.",
    "Chuẩn hóa error handler cấp ứng dụng để route không phải lặp try/catch và sendError.",
]: add_bullet(doc, item)

add_heading(doc, "7.1.3.2. Coupling của Backend", 3)
add_label(doc, "Ưu điểm")
for item in [
    "Router được mount qua apiModules; app.js chỉ phụ thuộc vào public contract path/router của module.",
    "Xác thực và phân quyền được tái sử dụng qua authMiddleware và requireRole, tránh để mỗi route tự giải mã token.",
    "Cấu hình cổng, client URL và reminder interval đã được gom trong config/env.js; CORS lấy dữ liệu từ module cấu hình thay vì đọc biến môi trường rải rác.",
    "Supabase client, Prisma client và JWT helper được tập trung trong lib, giúp caller không tự khởi tạo kết nối hay thuật toán token.",
]: add_bullet(doc, item)

add_label(doc, "Hạn chế")
for item in [
    "Hầu hết service import trực tiếp Supabase client và gọi tên bảng/cột. Business logic vì vậy phụ thuộc chặt vào SDK, schema và cách biểu diễn dữ liệu của Supabase.",
    "Nhiều service phụ thuộc chéo: appointment/boarding/staff gọi email service; doctor và email gọi settings service; admin gọi doctor schedule service. Những phụ thuộc này hợp lý về nghiệp vụ nhưng chưa có contract hoặc event boundary nên dễ tạo chuỗi thay đổi.",
    "Một số module đọc process.env trực tiếp ngoài config/env.js, ví dụ SMTP, clinic settings, timezone và storage bucket. Cấu hình chưa có một nguồn duy nhất.",
    "Response DTO chứa chi tiết trình bày frontend, làm backend gắn với Tailwind/theme và khó tái sử dụng cho mobile client hoặc API công khai.",
]: add_bullet(doc, item)

add_label(doc, "Đề xuất giảm Coupling")
for item in [
    "Đưa truy cập dữ liệu vào repository theo domain: AppointmentRepository, PetRepository, InvoiceRepository và BoardingRepository. Service phụ thuộc vào contract repository thay vì Supabase SDK.",
    "Dùng dependency injection ở composition root để truyền repository, email sender, clock và storage provider; unit test có thể thay bằng fake implementation.",
    "Dùng domain event hoặc application event cho email/thông báo sau các sự kiện AppointmentConfirmed, ExamCompleted, InvoicePaid và BoardingUpdated.",
    "Đưa toàn bộ biến môi trường vào config module có validation khi khởi động.",
]: add_bullet(doc, item)

add_heading(doc, "7.1.4. Kết luận về Coupling và Cohesion", 2)
add_body(doc, "Sau đợt tái cấu trúc, Pet Service có nền tảng kiến trúc tốt hơn đáng kể: frontend được tổ chức theo vai trò và feature; backend được tổ chức theo domain, có composition root, config, middleware và ranh giới app/server rõ ràng. Đây là bằng chứng của cohesion tốt ở cấp thư mục và coupling hợp lý ở cấp ứng dụng.")
add_body(doc, "Điểm cần ưu tiên không còn là 'tạo cấu trúc module từ đầu' mà là làm cho các ranh giới hiện có trở thành ranh giới thực sự. Cụ thể, cần chia nhỏ các God File, thống nhất HTTP client, đưa truy cập Supabase sau repository và tách view model khỏi backend. Khi đó cấu trúc tốt ở cấp thư mục sẽ được củng cố bằng cấu trúc tốt bên trong từng module.")

add_heading(doc, "7.2. Áp dụng các nguyên lý SOLID", 1)
add_heading(doc, "7.2.1. Giới thiệu chung", 2)
add_body(doc, "SOLID là tập hợp năm nguyên lý thiết kế giúp mã nguồn dễ hiểu, dễ kiểm thử, dễ bảo trì và mở rộng. Dù Pet Service sử dụng JavaScript/TypeScript theo phong cách module và functional composition nhiều hơn class inheritance, tinh thần SOLID vẫn áp dụng được cho function, component, service, module và contract dữ liệu.")
for item in [
    "S - Single Responsibility Principle (SRP): một module nên có một nhóm lý do thay đổi có liên quan chặt chẽ.",
    "O - Open/Closed Principle (OCP): thành phần nên mở cho mở rộng nhưng hạn chế sửa đổi phần ổn định.",
    "L - Liskov Substitution Principle (LSP): các implementation cùng contract phải thay thế được cho nhau mà không phá vỡ kỳ vọng của caller.",
    "I - Interface Segregation Principle (ISP): caller chỉ nên phụ thuộc vào contract nhỏ đúng nhu cầu.",
    "D - Dependency Inversion Principle (DIP): logic cấp cao nên phụ thuộc abstraction, không phụ thuộc trực tiếp chi tiết hạ tầng.",
]: add_bullet(doc, item)

add_heading(doc, "7.2.2. Đánh giá mức độ đáp ứng SOLID", 2)
add_heading(doc, "7.2.2.1. Single Responsibility Principle (SRP)", 3)
add_body(doc, "Pet Service đáp ứng SRP khá tốt ở cấp kiến trúc nhưng chưa đồng đều ở cấp file. app.js, server.js, config/env.js, auth.middleware.js và requestJson.ts là những ví dụ có trách nhiệm tương đối rõ. Việc chia frontend theo feature và backend theo domain cũng giúp mỗi khu vực có lý do thay đổi riêng.")
add_body(doc, "Ngược lại, medical.routes.js, admin.service.js, staff.service.js và các component lớn đang chứa nhiều use case. Chúng thay đổi vì nhiều nguyên nhân khác nhau: schema dữ liệu, nghiệp vụ, định dạng response, UI state hoặc tích hợp email/storage. Vì vậy SRP mới được đáp ứng một phần.")
add_table(doc, ["Module", "Đánh giá SRP", "Hướng cải thiện"], [
    ["backend/src/app.js", "Tốt: tạo app, middleware chung, mount module và health/404.", "Giữ nguyên phạm vi; thêm global error handler nếu cần."],
    ["backend/src/server.js", "Tốt: mở listener và scheduler.", "Tách scheduler thành adapter nếu cần kiểm thử sâu."],
    ["backend/src/modules/medical/medical.routes.js", "Chưa tốt: route, query, mapping, nghiệp vụ cùng file.", "Tách controller, exam service, specialist-order service, repository."],
    ["frontend/src/components/staff/StaffBoardingTab.tsx", "Chưa tốt: danh sách, phòng, checkout, upload, modal và state.", "Tách view/hook/modal theo từng use case."],
], [4.0, 6.0, 6.0])

add_heading(doc, "7.2.2.2. Open/Closed Principle (OCP)", 3)
add_body(doc, "OCP được thể hiện rõ nhất ở cơ chế apiModules. Một router mới có thể được đăng ký bằng một entry path/router mà không cần viết thêm handler trong app.js. Các feature và component UI cũng cho phép thêm màn hình mới bằng composition.")
add_body(doc, "Tuy nhiên App.tsx vẫn phải sửa khi thêm role; bảng status/label thường là object đóng và cần sửa trực tiếp khi có trạng thái mới; nhiều service là object lớn với các method cụ thể. Ranh giới reports và grooming re-export từ service khác cũng khiến việc mở rộng domain mới vẫn phải chạm vào module cũ.")
add_label(doc, "Đề xuất")
for item in [
    "Tạo portal registry theo role và route config theo feature.",
    "Tạo handler registry/strategy cho các chuyển trạng thái nghiệp vụ có thể mở rộng.",
    "Cho từng domain sở hữu service và mapper; module index chỉ công bố public API.",
    "Dùng enum/schema tập trung cho status thay vì switch và object rải rác.",
]: add_bullet(doc, item)

add_heading(doc, "7.2.2.3. Liskov Substitution Principle (LSP)", 3)
add_body(doc, "Dự án hiện không có hệ phân cấp class nghiệp vụ như Service -> HotelService/HealthService/SalonService trong ví dụ lý thuyết của báo cáo cũ. Vì vậy không nên khẳng định LSP dựa trên kế thừa. Cách đánh giá phù hợp hơn là khả năng thay thế implementation cùng contract.")
add_body(doc, "Một số dấu hiệu tích cực là response có trường ok/message tương đối nhất quán, service frontend trả về TypeScript type rõ ràng và các module index công bố public contract. Tuy nhiên chưa có interface repository/provider chính thức; Supabase client được import trực tiếp, nên chưa thể thay bằng Prisma repository hoặc fake repository mà không sửa service.")
add_label(doc, "Đề xuất")
for item in [
    "Định nghĩa contract Repository theo use case, bao gồm điều kiện lỗi và kiểu trả về; mọi implementation phải giữ cùng semantics.",
    "Chuẩn hóa ApiResult/ApiError và status code để caller không phải xử lý ngoại lệ theo từng endpoint.",
    "Viết contract test chạy cho SupabaseRepository và fake/in-memory repository để chứng minh khả năng thay thế.",
    "Chuẩn hóa DTO giữa snake_case từ database và camelCase ở API boundary.",
]: add_bullet(doc, item)

add_heading(doc, "7.2.2.4. Interface Segregation Principle (ISP)", 3)
add_body(doc, "TypeScript service và type theo customer domain là điểm mạnh của ISP: component chỉ import nhóm hàm và dữ liệu đúng nghiệp vụ. Backend module index cũng có thể đóng vai trò public facade thay vì caller import mọi chi tiết nội bộ.")
add_body(doc, "Hạn chế nằm ở các service object lớn. staffAppointmentsService chứa profile, appointment, grooming, boarding, payment và walk-in; doctorAppointmentsService chứa notification, schedule, exam và specialist order. Caller có thể import cùng object dù chỉ cần một nhóm nhỏ. Ở backend, admin.service.js phục vụ dashboard, user, service catalog, report, settings và lịch bác sĩ.")
add_label(doc, "Đề xuất")
for item in [
    "Tách StaffProfileApi, StaffAppointmentApi, GroomingApi, BoardingApi và PaymentApi.",
    "Tách DoctorScheduleApi, DoctorExamApi, DoctorNotificationApi và SpecialistOrderApi.",
    "Tách AdminDashboardService, AdminUserService, AdminCatalogService, AdminReportService và AdminSettingsService.",
    "Chỉ export contract công khai từ index; ẩn helper và chi tiết query trong module.",
]: add_bullet(doc, item)

add_heading(doc, "7.2.2.5. Dependency Inversion Principle (DIP)", 3)
add_body(doc, "Frontend đã có một lớp đảo phụ thuộc cơ bản: component gọi service thay vì truy cập database; backend route thường gọi service thay vì để UI hoặc request handler sở hữu toàn bộ nghiệp vụ. Các helper lib cũng che giấu phần khởi tạo JWT, Prisma và Supabase.")
add_body(doc, "Tuy nhiên business service vẫn phụ thuộc trực tiếp Supabase SDK, process.env, nodemailer/PDFKit hoặc service cụ thể khác. Đây là mức DIP chưa hoàn chỉnh: module cấp cao biết chi tiết hạ tầng cấp thấp và khó test độc lập.")
add_label(doc, "Đề xuất")
for item in [
    "Định nghĩa và inject AppointmentRepository, InvoiceRepository, EmailSender, FileStorage, PdfGenerator và Clock.",
    "Tạo composition root ở backend để ghép implementation thật; test ghép fake implementation.",
    "Tách domain event khỏi email side effect, tránh use case chính phụ thuộc trực tiếp SMTP.",
    "Bổ sung schema validation cho config và truyền config object thay vì đọc process.env trong service.",
]: add_bullet(doc, item)

add_heading(doc, "7.2.3. Bảng tổng hợp đánh giá", 2)
add_table(doc, ["Nguyên lý", "Mức đáp ứng", "Bằng chứng chính", "Ưu tiên"], [
    ["SRP", "Khá", "app/server tách rõ; domain module rõ; còn nhiều God File.", "Tách medical.routes, admin/staff service và component lớn."],
    ["OCP", "Khá", "apiModules và composition tốt; role/status còn hard-code.", "Registry/config cho portal và state transition."],
    ["LSP", "Trung bình", "Có TypeScript contract nhưng chưa có provider thay thế.", "Repository contract + contract test."],
    ["ISP", "Khá", "Customer API nhỏ; doctor/staff/admin service còn rộng.", "Tách API/service theo use case."],
    ["DIP", "Trung bình", "Có service boundary; business logic vẫn import Supabase trực tiếp.", "Repository + DI + domain event."],
], [2.0, 2.5, 6.0, 5.5])

add_heading(doc, "7.2.4. Lộ trình cải thiện đề xuất", 2)
add_label(doc, "Ngắn hạn")
for item in [
    "Thống nhất requestJson/ApiClient cho toàn bộ frontend và loại bỏ fetchWithAuth trùng lặp.",
    "Tách medical.routes.js thành router mỏng và service theo use case.",
    "Tách các component frontend trên 700-800 dòng thành hook, view và modal chuyên trách.",
    "Di chuyển class CSS/status color khỏi response backend về frontend design system.",
]: add_numbered(doc, item)
add_label(doc, "Trung hạn")
for item in [
    "Tách admin.service.js và staff.service.js theo domain/use case; chuyển ownership thật cho reports và grooming.",
    "Thiết lập repository layer cho appointment, pet, boarding và invoice trước, vì đây là các luồng có nhiều thao tác dữ liệu và side effect.",
    "Chuẩn hóa DTO, ApiResult và error middleware; bổ sung schema validation ở API boundary.",
]: add_numbered(doc, item)
add_label(doc, "Dài hạn")
for item in [
    "Áp dụng dependency injection tại composition root và contract test cho nhiều implementation.",
    "Dùng domain event cho email/notification, giảm phụ thuộc chéo giữa các service.",
    "Cân nhắc OpenAPI/shared schema để frontend-backend dùng cùng contract và hỗ trợ client mới.",
]: add_numbered(doc, item)

add_heading(doc, "7.2.5. Kết luận", 2)
add_body(doc, "Pet Service đã áp dụng tinh thần SOLID ở mức kiến trúc tốt hơn so với chương 7 ngắn trong báo cáo ban đầu. Điểm nổi bật là phân tách frontend/backend, tổ chức theo feature/domain, service layer, middleware dùng chung, composition root và ranh giới app/server. Các lựa chọn này giúp hệ thống phù hợp với tập nghiệp vụ lớn gồm đặt lịch, khám bệnh, grooming, boarding, hóa đơn và thông báo.")
add_body(doc, "Mức đáp ứng chưa đồng đều ở cấp triển khai. Những file rất lớn, truy cập Supabase trực tiếp, service object rộng và response chứa chi tiết UI là các dấu hiệu cần xử lý tiếp. Hướng cải thiện phù hợp nhất là củng cố ranh giới đã có bằng use-case service nhỏ, repository contract, dependency injection, domain event và contract dữ liệu thống nhất. Đây là lộ trình thực tế, không yêu cầu viết lại hệ thống nhưng tạo nền tảng vững cho bảo trì và mở rộng.")

add_heading(doc, "7.3. Các tệp mã nguồn tiêu biểu được đối chiếu", 1)
add_table(doc, ["Khu vực", "Tệp tiêu biểu", "Vai trò trong đánh giá"], [
    ["Backend composition", "backend/src/app.js; backend/src/server.js; backend/src/modules/index.js", "SRP, OCP và ranh giới khởi tạo ứng dụng."],
    ["Backend cross-cutting", "config/env.js; middlewares/auth.middleware.js; lib/supabaseClient.js; utils/http.js", "Cohesion, tái sử dụng và DIP."],
    ["Backend domain", "modules/medical, users, staff, appointments, boarding, notifications", "Cohesion nội bộ, coupling chéo và quy mô service."],
    ["Frontend shell", "frontend/src/app/App.tsx", "Session, role routing và coupling với portal."],
    ["Frontend services", "utils/requestJson.ts; services/customer/*; features/doctor/services/*; features/staff/services/*", "HTTP abstraction, DTO, ISP và duplication."],
    ["Frontend UI", "components/customer, doctor, staff, admin, ui", "Phân chia component, cohesion và God Component."],
], [3.2, 7.3, 5.5])

add_callout(doc, "Ghi chú", "Đánh giá phản ánh mã nguồn hiện tại, gồm backend đã tái cấu trúc nhưng chưa commit; số dòng chỉ dùng để nhận diện điểm nóng.")

doc.core_properties.title = "Chương 7 - Nguyên tắc thiết kế Pet Service"
doc.core_properties.subject = "Coupling, Cohesion và SOLID"
doc.core_properties.author = "Nhóm phát triển Pet Service"
doc.core_properties.keywords = "Pet Service, ITSS, SOLID, Coupling, Cohesion"
doc.save(OUTPUT)
print(OUTPUT)
