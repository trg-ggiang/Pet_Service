from pathlib import Path
import re
import zipfile

from docx import Document
from pypdf import PdfReader


DOCX = Path(r"D:\Pet_Sevice\output\documents\Chuong_7_Design_Concepts_SOLID_Pet_Service_70_30.docx")
PDF = Path(r"D:\Pet_Sevice\tmp\chapter7-rebalanced-render\chapter7.pdf")

doc = Document(DOCX)
text = "\n".join(p.text for p in doc.paragraphs)
text += "\n" + "\n".join(
    cell.text for table in doc.tables for row in table.rows for cell in row.cells
)

with zipfile.ZipFile(DOCX) as archive:
    xml = "\n".join(
        archive.read(name).decode("utf-8", "ignore")
        for name in archive.namelist()
        if name == "word/document.xml"
        or name.startswith("word/header")
        or name.startswith("word/footer")
    )

colors = set(re.findall(r'w:color[^>]*w:val="([^"]+)"', xml))
fills = set(re.findall(r'w:shd[^>]*w:fill="([^"]+)"', xml))
pages = len(PdfReader(PDF).pages)
stats = {
    "pages": pages,
    "characters": len(text),
    "replacement_characters": text.count(chr(0xFFFD)),
    "text_colors": sorted(colors),
    "cell_fills": sorted(fills),
}
print(stats)

assert pages == 8
assert text.count(chr(0xFFFD)) == 0
assert colors <= {"000000", "auto"}
assert fills <= {"FFFFFF", "auto"}
assert "7.1. Áp dụng Design Concepts" in text
assert "7.2. Áp dụng Design Principles SOLID" in text
assert "Những điểm đã làm tốt" in text
assert "Nhóm 23" in text
assert all(name in text for name in ("Trần Đức Nam Anh", "Nguyễn Đức Hiếu", "Triệu Trường Giang"))
