from docx import Document
from pypdf import PdfReader
from pathlib import Path
import re
import zipfile

docx_path = Path(r"D:\Pet_Sevice\output\documents\Chuong_7_Good_Design_Pet_Service_Nhom_23_Den_Trang.docx")
pdf_path = Path(r"D:\Pet_Sevice\tmp\good-design-render\chapter7-black-white.pdf")
doc = Document(docx_path)
text = "\n".join(p.text for p in doc.paragraphs)
text += "\n" + "\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
pages = len(PdfReader(pdf_path).pages)
with zipfile.ZipFile(docx_path) as archive:
    xml = "\n".join(
        archive.read(name).decode("utf-8", "ignore")
        for name in archive.namelist()
        if name == "word/document.xml"
        or name.startswith("word/header")
        or name.startswith("word/footer")
    )
colors = set(re.findall(r'w:color[^>]*w:val="([^"]+)"', xml))
fills = set(re.findall(r'w:shd[^>]*w:fill="([^"]+)"', xml))
print({"pages": pages, "chars": len(text), "replacement": text.count(chr(0xFFFD)), "colors": sorted(colors), "fills": sorted(fills)})
assert pages == 15
assert "CHƯƠNG 7. TÀI LIỆU THIẾT KẾ TỐT" in text
assert "7.4.3.5. Dependency Inversion Principle" in text
assert colors <= {"000000", "auto"}
assert fills <= {"FFFFFF", "auto"}
assert all(name in text for name in ("Trần Đức Nam Anh", "Nguyễn Đức Hiếu", "Triệu Trường Giang"))
