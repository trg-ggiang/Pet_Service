from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = r"D:\Pet_Sevice\output\documents\Chuong_7_Good_Design_Pet_Service_Nhom_23_Den_Trang.docx"
FONT = "Times New Roman"
MONO = "Consolas"
BLACK = "000000"
NAVY = "000000"
BLUE = "000000"
LIGHT_BLUE = "FFFFFF"
LIGHT_GRAY = "FFFFFF"
MID_GRAY = "FFFFFF"
MUTED = "000000"
WHITE = "FFFFFF"


def set_font(run, size=12.5, bold=False, italic=False, color=BLACK, name=FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    rpr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(p, before=0, after=5, line=1.25, align=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        p.alignment = align


def body(doc, text, indent=True, after=5):
    p = doc.add_paragraph()
    set_spacing(p, after=after, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.9)
    set_font(p.add_run(text))
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    size = {1: 16, 2: 14, 3: 12.5}[level]
    before = {1: 10, 2: 8, 3: 5}[level]
    after = {1: 7, 2: 5, 3: 3}[level]
    set_spacing(p, before=before, after=after, line=1.1)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.outline_level = level - 1
    set_font(p.add_run(text), size=size, bold=True, color=NAVY if level < 3 else BLACK)
    return p


def label(doc, text, color=BLUE):
    p = doc.add_paragraph()
    set_spacing(p, before=3, after=2, line=1.15)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), size=12.5, bold=True, color=color)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    set_spacing(p, after=3, line=1.2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.left_indent = Cm(0.72 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    if p.runs:
        p.runs[0].text = text
        set_font(p.runs[0], size=12.2)
    else:
        set_font(p.add_run(text), size=12.2)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_spacing(p, after=3, line=1.2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    if p.runs:
        p.runs[0].text = text
        set_font(p.runs[0], size=12.2)
    else:
        set_font(p.add_run(text), size=12.2)
    return p


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def no_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:cantSplit")
    tr_pr.append(element)


def table(doc, headers, rows, widths_cm, font_size=10.2):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    repeat_header(t.rows[0])
    no_split(t.rows[0])
    for i, text in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = Cm(widths_cm[i])
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_fill(c, WHITE)
        set_cell_margins(c)
        p = c.paragraphs[0]
        set_spacing(p, after=0, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_font(p.add_run(text), size=font_size, bold=True, color=BLACK)
    for ridx, values in enumerate(rows):
        row = t.add_row()
        no_split(row)
        for i, value in enumerate(values):
            c = row.cells[i]
            c.width = Cm(widths_cm[i])
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(c)
            if ridx % 2:
                set_cell_fill(c, LIGHT_GRAY)
            p = c.paragraphs[0]
            set_spacing(p, after=0, line=1.1, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_font(p.add_run(str(value)), size=font_size)
    spacer = doc.add_paragraph()
    set_spacing(spacer, after=1, line=1)
    return t


def callout(doc, title, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    row = t.rows[0]
    no_split(row)
    c = row.cells[0]
    c.width = Cm(15.8)
    set_cell_fill(c, LIGHT_BLUE)
    set_cell_margins(c, top=130, start=160, bottom=130, end=160)
    p = c.paragraphs[0]
    set_spacing(p, after=0, line=1.18, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    set_font(p.add_run(title + ": "), size=11.5, bold=True, color=NAVY)
    set_font(p.add_run(text), size=11.5)
    spacer = doc.add_paragraph()
    set_spacing(spacer, after=1, line=1)


def code_block(doc, code):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    row = t.rows[0]
    no_split(row)
    c = row.cells[0]
    c.width = Cm(15.5)
    set_cell_fill(c, WHITE)
    set_cell_margins(c, top=130, start=170, bottom=130, end=170)
    p = c.paragraphs[0]
    set_spacing(p, after=0, line=1.08)
    for idx, line in enumerate(code.splitlines()):
        if idx:
            p.add_run().add_break()
        set_font(p.add_run(line), size=9.4, color=BLACK, name=MONO)
    spacer = doc.add_paragraph()
    set_spacing(spacer, after=1, line=1)


def page_break(doc):
    doc.add_page_break()


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_font(run, size=9.5, color=MUTED)


def toc_row(doc, title, page, level=0):
    p = doc.add_paragraph()
    set_spacing(p, after=3, line=1.15)
    p.paragraph_format.left_indent = Cm(level * 0.55)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15.5))
    set_font(p.add_run(title), size=11.8, bold=(level == 0))
    set_font(p.add_run("\t" + str(page)), size=11.8, bold=(level == 0))


doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.1)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(3.0)
sec.right_margin = Cm(2.0)
sec.header_distance = Cm(1.0)
sec.footer_distance = Cm(1.0)
sec.different_first_page_header_footer = True

# Academic-report override of narrative_proposal: A4, Times New Roman, 12.5 pt.
normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(12.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.25
for style_name in ("List Bullet", "List Bullet 2", "List Number"):
    style = doc.styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(12.2)

header = sec.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_spacing(hp, after=0, line=1)
set_font(hp.add_run("CHƯƠNG 7  |  GOOD DESIGN SPECIFICATION  |  PET SERVICE  |  NHÓM 23"), size=8.8, color=BLACK)
footer = sec.footer
set_spacing(footer.paragraphs[0], after=0, line=1)
add_page_field(footer.paragraphs[0])

# PAGE 1 - COVER
p = doc.add_paragraph()
set_spacing(p, before=0, after=3, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(p.add_run("ĐẠI HỌC BÁCH KHOA HÀ NỘI"), size=13, bold=True)
p = doc.add_paragraph()
set_spacing(p, after=40, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(p.add_run("TRƯỜNG CÔNG NGHỆ THÔNG TIN VÀ TRUYỀN THÔNG"), size=13, bold=True)
p = doc.add_paragraph()
set_spacing(p, after=6, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(p.add_run("CHƯƠNG 7. TÀI LIỆU THIẾT KẾ TỐT"), size=20, bold=True, color=BLACK)
p = doc.add_paragraph()
set_spacing(p, after=8, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(p.add_run("(Good Design Specification)"), size=13, italic=True, color=MUTED)
p = doc.add_paragraph()
set_spacing(p, after=28, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(p.add_run("HỆ THỐNG QUẢN LÝ TRUNG TÂM CHĂM SÓC THÚ CƯNG\nPET SERVICE"), size=16, bold=True)
p = doc.add_paragraph()
set_spacing(p, after=4, line=1.2, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(p.add_run("Học phần: Phát triển phần mềm theo chuẩn kỹ năng ITSS"), size=12.5)
p = doc.add_paragraph()
set_spacing(p, after=18, line=1.2, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(p.add_run("GVHD: ThS. Nguyễn Mạnh Tuấn"), size=12.5)
p = doc.add_paragraph()
set_spacing(p, after=5, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(p.add_run("NHÓM 23"), size=14, bold=True, color=NAVY)
table(doc, ["STT", "Họ và tên sinh viên"], [
    ["1", "Trần Đức Nam Anh"],
    ["2", "Nguyễn Đức Hiếu"],
    ["3", "Triệu Trường Giang"],
], [2.3, 10.5], font_size=11.2)
p = doc.add_paragraph()
set_spacing(p, before=32, after=0, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(p.add_run("Hà Nội, tháng 6 năm 2026"), size=12)

# PAGE 2 - TOC
page_break(doc)
heading(doc, "MỤC LỤC", 1)
for title, pg, level in [
    ("7.1. Thiết kế mô-đun (Modularity)", 3, 0),
    ("7.1.1. Khái niệm và vai trò", 3, 1),
    ("7.1.2. Đặc điểm của một thiết kế mô-đun tốt", 3, 1),
    ("7.2. Coupling và Cohesion", 4, 0),
    ("7.2.1. Các mức Coupling", 4, 1),
    ("7.2.2. Các mức Cohesion", 4, 1),
    ("7.3. Tổng quan các nguyên lý SOLID", 4, 0),
    ("7.4. Áp dụng Good Design vào dự án Pet Service", 5, 0),
    ("7.4.1. Thiết kế mô-đun của hệ thống", 5, 1),
    ("7.4.1.1. Kiến trúc tổng thể", 5, 2),
    ("7.4.1.2. Modularity phía Backend", 6, 2),
    ("7.4.1.3. Modularity phía Frontend", 7, 2),
    ("7.4.2. Đánh giá Coupling và Cohesion", 8, 1),
    ("7.4.3. Đánh giá các nguyên lý SOLID", 10, 1),
    ("7.5. Đánh giá theo các tính chất mô-đun lý tưởng", 12, 0),
    ("7.6. Các điểm cần tiếp tục hoàn thiện", 13, 0),
    ("7.7. Lộ trình cải thiện thiết kế", 14, 0),
    ("7.8. Kết luận", 15, 0),
]: toc_row(doc, title, pg, level)

# PAGE 3 - MODULARITY THEORY
page_break(doc)
heading(doc, "7.1. Thiết kế mô-đun (Modularity)", 1)
heading(doc, "7.1.1. Khái niệm và vai trò", 2)
body(doc, "Mô-đun là một đơn vị tương đối độc lập trong hệ thống phần mềm, có chức năng, dữ liệu đầu vào, dữ liệu đầu ra và phạm vi trách nhiệm rõ ràng. Mô-đun có thể là một component, một service, một route group, một tập hợp type hoặc một thư mục nghiệp vụ. Thiết kế mô-đun không đơn giản là chia code thành nhiều file; điều quan trọng hơn là mỗi phần được chia ra phải có ý nghĩa và có ranh giới giao tiếp rõ.")
body(doc, "Khi hệ thống được chia mô-đun hợp lý, thành viên trong nhóm có thể đọc, sửa và kiểm thử một phần mà không cần hiểu toàn bộ ứng dụng. Điều này đặc biệt cần thiết với Pet Service vì dự án có nhiều tác nhân cùng sử dụng: khách hàng, nhân viên, bác sĩ và quản trị viên. Mỗi vai trò có màn hình, dữ liệu và quy trình nghiệp vụ riêng nhưng vẫn phải phối hợp với nhau trong các luồng đặt lịch, khám bệnh, grooming, lưu trú và thanh toán.")
label(doc, "Lợi ích chính của tính mô-đun:")
for x in [
    "Dễ hiểu: cấu trúc thư mục thể hiện được các vùng chức năng của hệ thống.",
    "Dễ phân công: thành viên có thể làm việc trên các module khác nhau và giảm xung đột mã nguồn.",
    "Dễ kiểm thử: lỗi được khoanh vùng trong một nhóm chức năng nhỏ hơn.",
    "Dễ bảo trì: thay đổi trong một nghiệp vụ ít làm ảnh hưởng sang nghiệp vụ khác.",
    "Dễ tái sử dụng: component UI, helper và service có thể dùng lại ở nhiều màn hình.",
]: bullet(doc, x)
heading(doc, "7.1.2. Đặc điểm của một thiết kế mô-đun tốt", 2)
table(doc, ["Tính chất", "Ý nghĩa trong thiết kế phần mềm"], [
    ["Decomposable", "Có thể phân rã hệ thống thành các phần nhỏ để giảm độ phức tạp."],
    ["Composable", "Các mô-đun có thể lắp ghép và phối hợp thành một luồng hoàn chỉnh."],
    ["Understandable", "Có thể hiểu và kiểm thử một mô-đun mà không phải đọc quá nhiều phần khác."],
    ["Continuity", "Một thay đổi nhỏ trong yêu cầu chỉ kéo theo thay đổi ở số ít mô-đun."],
    ["Isolation", "Lỗi hoặc thay đổi hạ tầng được giới hạn, không lan truyền toàn hệ thống."],
], [4.0, 12.0], font_size=10.6)
callout(doc, "Nhận xét", "Thiết kế tốt là sự cân bằng. Chia quá ít sẽ tạo file lớn và khó bảo trì; chia quá nhỏ lại làm tăng số lượng phụ thuộc và khiến luồng xử lý bị phân mảnh.")

# PAGE 4 - COUPLING, COHESION, SOLID THEORY
page_break(doc)
heading(doc, "7.2. Coupling và Cohesion", 1)
body(doc, "Mục tiêu quan trọng của thiết kế phần mềm là đạt High Cohesion và Low Coupling: các thành phần trong cùng một mô-đun liên quan chặt chẽ với nhau, trong khi sự phụ thuộc giữa các mô-đun được giữ ở mức cần thiết và thông qua contract rõ ràng.")
heading(doc, "7.2.1. Các mức Coupling", 2)
table(doc, ["Mức độ", "Mô tả", "Đánh giá"], [
    ["Content", "Một mô-đun can thiệp trực tiếp vào chi tiết nội bộ của mô-đun khác.", "Xấu nhất"],
    ["Common", "Nhiều mô-đun dùng chung trạng thái hoặc dữ liệu toàn cục.", "Rất kém"],
    ["Control", "Truyền cờ điều khiển để quyết định luồng xử lý bên trong mô-đun khác.", "Kém"],
    ["Stamp", "Truyền cả cấu trúc lớn nhưng bên nhận chỉ dùng một phần nhỏ.", "Trung bình"],
    ["Data", "Chỉ truyền dữ liệu cần thiết thông qua tham số hoặc DTO rõ ràng.", "Tốt nhất"],
], [3.0, 10.2, 2.8], font_size=9.8)
heading(doc, "7.2.2. Các mức Cohesion", 2)
table(doc, ["Mức độ", "Mô tả", "Đánh giá"], [
    ["Coincidental", "Các phần tử được gom ngẫu nhiên, gần như không liên quan.", "Thấp"],
    ["Logical", "Nhóm theo một ý nghĩa chung nhưng nhiệm vụ cụ thể khác nhau.", "Rất kém"],
    ["Temporal", "Nhóm vì được thực hiện cùng thời điểm.", "Kém"],
    ["Procedural", "Nhóm vì được gọi theo một trình tự xử lý.", "Trung bình"],
    ["Communicational", "Các phần cùng thao tác trên một tập dữ liệu chung.", "Khá"],
    ["Functional", "Mọi phần phối hợp để hoàn thành một trách nhiệm rõ ràng.", "Tốt nhất"],
], [3.0, 10.2, 2.8], font_size=9.6)
heading(doc, "7.3. Tổng quan các nguyên lý SOLID", 1)
body(doc, "SOLID gồm SRP, OCP, LSP, ISP và DIP. Các nguyên lý này không chỉ dành cho class hướng đối tượng mà còn có thể áp dụng cho component React, module JavaScript, service, middleware và contract TypeScript.", indent=False, after=2)

# PAGE 5 - APPLICATION OVERVIEW
page_break(doc)
heading(doc, "7.4. Áp dụng Good Design vào dự án Pet Service", 1)
body(doc, "Pet Service là ứng dụng fullstack quản lý trung tâm chăm sóc thú cưng. Frontend sử dụng React 18, Vite và TypeScript; backend sử dụng Express, Prisma/PostgreSQL và Supabase server client. Hệ thống hỗ trợ bốn vai trò chính: customer, staff, doctor và admin. Các luồng nghiệp vụ bao gồm quản lý hồ sơ thú cưng, đặt lịch, khám bệnh, kê đơn, grooming, lưu trú, hóa đơn, thanh toán, báo cáo và thông báo.")
callout(doc, "Phạm vi đánh giá", "Tài liệu phản ánh code hiện tại trong thư mục Pet_Sevice, bao gồm backend đã được tái cấu trúc theo domain module nhưng chưa commit. Trọng tâm là ghi nhận các nền tảng thiết kế tốt mà nhóm đã xây dựng; các đề xuất còn lại được trình bày như bước hoàn thiện tiếp theo.")
heading(doc, "7.4.1. Thiết kế mô-đun của hệ thống", 2)
heading(doc, "7.4.1.1. Kiến trúc tổng thể", 3)
body(doc, "Ở mức cao nhất, dự án tách frontend và backend thành hai ứng dụng độc lập. Frontend chịu trách nhiệm hiển thị, quản lý trạng thái giao diện và gọi HTTP API. Backend chịu trách nhiệm xác thực, phân quyền, kiểm tra nghiệp vụ và truy cập dữ liệu. PostgreSQL/Supabase nằm phía sau backend nên trình duyệt không truy vấn database trực tiếp.")
code_block(doc, "React/Vite Frontend\n        |  HTTP + JSON + JWT\n        v\nExpress Backend\n        |  Supabase server client / Prisma\n        v\nPostgreSQL Database")
label(doc, "Luồng ví dụ - đặt lịch khám:")
for x in [
    "CustomerAppointmentsTab mở form và gọi customerAppointmentsApi.ts.",
    "Request được gửi đến /api/customer/appointments kèm token.",
    "authMiddleware xác thực người dùng và requireRole kiểm tra vai trò customer.",
    "customer.routes.js chuyển dữ liệu sang customer-appointment.service.js.",
    "Service kiểm tra thú cưng, thời gian, bác sĩ hoặc nhân viên rồi ghi dữ liệu qua Supabase.",
]: numbered(doc, x)
body(doc, "Luồng trên cho thấy mỗi tầng có nhiệm vụ tương đối khác nhau. Đây là nền tảng của Separation of Concerns và giúp dự án có thể tiếp tục tách sâu hơn mà không phải viết lại toàn bộ hệ thống.")

# PAGE 6 - BACKEND MODULARITY
page_break(doc)
heading(doc, "7.4.1.2. Modularity phía Backend", 2)
body(doc, "Đợt tái cấu trúc hiện tại đã chuyển backend từ cấu trúc routes/services rời rạc sang cấu trúc domain module tại backend/src/modules. Các module đang có gồm auth, users, customers, pets, services, appointments, staff, doctors, medical, grooming, boarding, billing, notifications và reports. Số lượng và mức hoàn thiện của các module được ghi theo source thực tế, không sử dụng con số 22 module như tài liệu tham khảo.")
table(doc, ["Nhóm", "Module tiêu biểu", "Trách nhiệm"], [
    ["Nền tảng", "auth, users", "Đăng nhập, đăng ký, khôi phục mật khẩu, tài khoản và cài đặt."],
    ["Khách hàng", "customers, pets, appointments", "Hồ sơ, thú cưng, lịch hẹn và lịch sử dịch vụ."],
    ["Chuyên môn", "doctors, medical", "Lịch bác sĩ, ca khám, bệnh án, đơn thuốc và chuyên khoa."],
    ["Vận hành", "staff, grooming, boarding", "Tiếp nhận, làm đẹp, lưu trú và chăm sóc hằng ngày."],
    ["Tài chính", "billing, reports", "Hóa đơn PDF, doanh thu và số liệu báo cáo."],
    ["Liên lạc", "notifications", "Thông báo trong hệ thống, email và lịch nhắc."],
], [3.0, 5.0, 8.0], font_size=10.0)
label(doc, "Các ranh giới kỹ thuật đã được tách rõ:")
for x in [
    "app.js tạo Express app, cài middleware chung và mount các domain router.",
    "server.js chỉ mở cổng và bật reminder scheduler khi chạy trực tiếp.",
    "config chứa env và CORS; middlewares chứa xác thực/phân quyền.",
    "lib chứa JWT, Prisma và Supabase client; utils chứa helper HTTP.",
    "modules/index.js đóng vai trò nơi đăng ký prefix và router công khai.",
]: bullet(doc, x)
code_block(doc, "for (const apiModule of apiModules) {\n  app.use(apiModule.path, apiModule.router);\n}")
body(doc, "Cách mount module bằng danh sách giúp app.js không cần biết từng endpoint con. Khi thêm một domain router mới, nhóm chỉ cần đăng ký path và router trong modules/index.js. Đây là một ví dụ khá rõ của modularity và OCP ở cấp ứng dụng.")

# PAGE 7 - FRONTEND MODULARITY
page_break(doc)
heading(doc, "7.4.1.3. Modularity phía Frontend", 2)
body(doc, "Frontend được tổ chức theo hai hướng song song: chia theo vai trò ở thư mục features và chia theo loại thành phần ở components, services, types, utils. Cách tổ chức này phù hợp với React vì page có thể làm container, component chịu trách nhiệm hiển thị, service gọi API và type mô tả contract dữ liệu.")
table(doc, ["Khu vực", "Ví dụ", "Vai trò"], [
    ["features/auth", "LoginPage, RegisterPage", "Điều phối form và luồng xác thực."],
    ["features/customer", "CustomerPortal, CustomerPetsPage", "Màn hình và điều hướng của khách hàng."],
    ["features/doctor", "DoctorPortal, DoctorExamScreen", "Lịch khám, bệnh án, thống kê và khám bệnh."],
    ["features/staff", "StaffPortal", "Tiếp nhận, grooming, boarding và thanh toán."],
    ["features/admin", "Dashboard, Users, Reports, Settings", "Quản trị toàn hệ thống."],
    ["components/ui", "Button, Dialog, Table, Input", "Thành phần trình bày dùng chung."],
], [3.2, 6.0, 6.8], font_size=10.1)
label(doc, "Ví dụ về contract API phía customer:")
code_block(doc, "export async function requestJson<T>(url: string, init?: RequestInit): Promise<T> {\n  const response = await fetch(apiUrl(url), { ...init, headers: { ... } });\n  const payload = await response.json().catch(() => ({}));\n  if (!response.ok || payload?.ok === false) throw new Error(payload?.message);\n  return payload as T;\n}")
body(doc, "requestJson.ts gom việc xác định URL, gắn header, parse JSON và chuyển lỗi HTTP thành Error. customerAppointmentsApi.ts và các service customer khác dùng helper này, vì vậy component không phải lặp lại phần kỹ thuật của fetch.")
label(doc, "Kết quả đạt được và phần tiếp tục chuẩn hóa:")
for x in [
    "Điểm tốt: feature theo role rõ, customer API và type được chia theo nghiệp vụ, component UI có khả năng tái sử dụng.",
    "Các role portal có không gian riêng nên việc phát triển song song và kiểm tra quyền truy cập thuận lợi hơn.",
    "Hệ thống type giúp nhiều luồng appointment, medication, notification và pet có contract rõ ràng.",
    "Phần tiếp tục chuẩn hóa là hợp nhất fetchWithAuth/requestJson và tách dần một số component dài.",
]: bullet(doc, x)

# PAGE 8 - COUPLING APPLICATION
page_break(doc)
heading(doc, "7.4.2. Đánh giá Coupling và Cohesion", 1)
heading(doc, "7.4.2.1. Coupling trong Pet Service", 2)
body(doc, "Coupling của hệ thống được đánh giá ở hai chiều: coupling giữa frontend với backend và coupling giữa các module trong backend. Kiến trúc hiện tại đã tránh Content Coupling ở cấp frontend vì trình duyệt không can thiệp trực tiếp vào dữ liệu nội bộ của backend. Hai phía giao tiếp qua HTTP, JSON và token.")
label(doc, "Các điểm thể hiện coupling tương đối thấp:")
for x in [
    "Frontend gọi service API thay vì gọi Supabase SDK hoặc biết câu truy vấn database.",
    "App mount module qua public contract path/router; chi tiết route nằm trong domain.",
    "Xác thực và phân quyền dùng authMiddleware/requireRole thay vì lặp ở từng endpoint.",
    "JWT, Prisma và Supabase client được tạo tập trung trong lib.",
    "Service frontend trả DTO/type thay vì chuyển nguyên row database vào component.",
    "Các route group có prefix riêng theo vai trò và nghiệp vụ, giúp thay đổi một nhóm API ít tác động sang nhóm khác.",
    "Compatibility re-export ở frontend giữ import cũ hoạt động trong quá trình tái cấu trúc.",
]: bullet(doc, x)
label(doc, "Các điểm có thể giảm coupling thêm:")
for x in [
    "Nhiều business service import Supabase client trực tiếp và gọi tên bảng/cột. Khi đổi data provider, nhiều file phải sửa.",
    "Một số response backend chứa class Tailwind như cls, dot hoặc mã màu. Backend vì vậy biết chi tiết hiển thị của frontend.",
    "Các cách gọi HTTP ở frontend chưa thống nhất, làm cơ chế token và error handling bị lặp.",
]: bullet(doc, x)
table(doc, ["Quan hệ", "Mức hiện tại", "Hướng tốt hơn"], [
    ["UI - API", "Data/Stamp coupling qua DTO", "Giữ DTO nhỏ, ổn định và có schema."],
    ["Route - Service", "Khá lỏng ở auth/customer, chặt hơn ở medical", "Controller/route mỏng, không query trực tiếp."],
    ["Service - Database", "Coupling trực tiếp với Supabase", "Repository contract hoặc data-access module."],
    ["Service - Email", "Import và gọi trực tiếp", "EmailSender/event publisher."],
], [4.0, 6.0, 6.0], font_size=10.1)
callout(doc, "Kết luận Coupling", "Phần lớn giao tiếp chính đã đi qua API, middleware, service và public module contract nên coupling ở cấp kiến trúc được kiểm soát khá tốt. Phần cần hoàn thiện tập trung ở data-access và việc chuẩn hóa HTTP client.")

# PAGE 9 - COHESION APPLICATION
page_break(doc)
heading(doc, "7.4.2.2. Cohesion trong Pet Service", 2)
body(doc, "Ở cấp thư mục, backend có cohesion khá tốt vì module được đặt theo domain. Ở cấp file, mức cohesion không đồng đều. Một số file có phạm vi rất rõ, nhưng một số file lớn đang gom nhiều use case và nhiều loại công việc khác nhau.")
label(doc, "Ví dụ cohesion tốt:")
for x in [
    "config/env.js chỉ đọc và chuẩn hóa cấu hình cổng, client URL và reminder interval.",
    "utils/http.js chỉ chứa helper sendError dùng để trả lỗi HTTP.",
    "app.js tập trung tạo ứng dụng Express; server.js tập trung khởi động process.",
    "customer-profile.service.js xử lý đọc/cập nhật hồ sơ customer.",
    "customer-medications.service.js tập trung lấy danh sách thuốc đang sử dụng.",
    "auth.routes.js và auth.service.js đã tách giao tiếp HTTP khỏi phần lớn logic xác thực.",
    "Các component doctor nhỏ như DoctorScheduleHeader, DoctorScheduleStats và VitalInput có trách nhiệm hiển thị rõ.",
]: bullet(doc, x)
label(doc, "Các phạm vi nên tiếp tục tách nhỏ:")
for x in [
    "medical.routes.js gần 2.000 dòng, vừa route, query, mapping, khám bệnh, đơn thuốc, chuyên khoa và notification.",
    "admin.service.js khoảng 1.700 dòng, chứa dashboard, user, dịch vụ, nhân sự, báo cáo, lịch bác sĩ và cài đặt.",
    "Một số component portal dài có thể tách thêm theo hook, view và modal để giữ cohesion ở cấp file.",
]: bullet(doc, x)
table(doc, ["Tệp", "Cohesion ước lượng", "Lý do"], [
    ["backend/src/app.js", "Functional", "Các phần cùng phục vụ việc tạo Express app."],
    ["backend/src/config/env.js", "Functional", "Chỉ xử lý cấu hình runtime."],
    ["medical.routes.js", "Procedural/Communicational", "Cùng luồng khám nhưng ôm nhiều trách nhiệm kỹ thuật."],
    ["admin.service.js", "Logical", "Nhiều chức năng admin được gom theo vai trò."],
    ["StaffBoardingTab.tsx", "Logical/Procedural", "Các thao tác boarding liên quan nhưng quá nhiều use case UI."],
], [5.4, 4.0, 6.6], font_size=9.9)
label(doc, "Hướng nâng cao cohesion:")
for x in [
    "Tách theo use case nhỏ: exam, prescription, specialist order, notification.",
    "Tách mapper/formatter khỏi service truy vấn và khỏi route.",
    "Tách component lớn thành container, hook, view và modal chuyên trách.",
]: bullet(doc, x)
callout(doc, "Kết luận Cohesion", "Cohesion ở cấp thư mục và nhiều file hạ tầng đã tốt sau tái cấu trúc. Một số file nghiệp vụ lớn là phần còn lại cần tách dần để Functional Cohesion được áp dụng đồng đều hơn.")

# PAGE 10 - SRP/OCP
page_break(doc)
heading(doc, "7.4.3. Đánh giá các nguyên lý SOLID", 1)
heading(doc, "7.4.3.1. Single Responsibility Principle (SRP)", 2)
body(doc, "SRP yêu cầu một thành phần chỉ nên có một nhóm lý do thay đổi. Trong Pet Service, ví dụ rõ nhất là app.js và server.js. Nếu thay cách mount route hoặc middleware chung thì sửa app.js; nếu thay cổng hoặc scheduler thì sửa server.js. Hai loại thay đổi không còn nằm trong cùng một file.")
code_block(doc, "// app.js\nconst app = express();\napp.use(cors(corsOptions));\nfor (const apiModule of apiModules) app.use(apiModule.path, apiModule.router);\n\n// server.js\nif (require.main === module) {\n  app.listen(env.port);\n  startReminderScheduler();\n}")
label(doc, "Các điểm tuân thủ SRP:")
for x in [
    "auth.middleware.js chỉ xử lý bearer token, user context và role.",
    "invoice-pdf.service.js tập trung xây dựng hóa đơn PDF.",
    "Các component nhỏ như DoctorScheduleHeader, DoctorScheduleStats và VitalInput có phạm vi hiển thị rõ.",
    "Customer profile, medications và notifications có service chuyên biệt theo từng nhu cầu.",
]: bullet(doc, x)
label(doc, "Phần cần hoàn thiện thêm:")
for x in [
    "Một số file nghiệp vụ lớn vẫn có nhiều lý do thay đổi và nên được chia theo use case.",
]: bullet(doc, x)
heading(doc, "7.4.3.2. Open-Closed Principle (OCP)", 2)
body(doc, "OCP yêu cầu thành phần mở để mở rộng nhưng hạn chế sửa code lõi đã ổn định. apiModules là ví dụ tích cực: có thể thêm một route group bằng entry mới. Các component UI cũng có thể được kết hợp để tạo page mới mà không sửa implementation của Button, Dialog hoặc Table.")
label(doc, "Các điểm đã hỗ trợ mở rộng:")
for x in [
    "apiModules cho phép đăng ký thêm domain router mà app.js không cần biết endpoint con.",
    "Component UI dùng composition, giúp page mới tái sử dụng Button, Card, Dialog, Table và form control.",
    "Feature service và compatibility re-export giúp di chuyển module mà hạn chế làm hỏng caller cũ.",
]: bullet(doc, x)
label(doc, "Phần có thể cải thiện:")
for x in [
    "App.tsx phải sửa khi thêm role mới vì portal được chọn bằng chuỗi if.",
    "Một số bảng status/label có thể chuyển sang registry hoặc policy để mở rộng thuận tiện hơn.",
]: bullet(doc, x)

# PAGE 11 - LSP/ISP
page_break(doc)
heading(doc, "7.4.3.3. Liskov Substitution Principle (LSP)", 2)
body(doc, "Pet Service chủ yếu dùng JavaScript/TypeScript theo phong cách module và composition, không có nhiều class cha - class con trong nghiệp vụ. Vì vậy, không nên cố chứng minh LSP bằng một hệ kế thừa không tồn tại. Cách phù hợp hơn là xem các implementation cùng contract có thay thế được cho nhau hay không.")
label(doc, "Nền tảng hiện tại:")
for x in [
    "Frontend có TypeScript interface mô tả DTO của appointment, medication, notification và pet.",
    "Các service thường trả dữ liệu theo shape ổn định và lỗi dưới dạng Error/message.",
    "Các module index tạo public contract ở cấp module và giúp caller không cần import mọi file nội bộ.",
    "Cách chuẩn hóa DTO hiện có là tiền đề để bổ sung repository contract trong bước tiếp theo.",
]: bullet(doc, x)
label(doc, "Điều kiện để đáp ứng LSP tốt hơn:")
for x in [
    "Định nghĩa contract cho repository: kiểu dữ liệu, điều kiện lỗi và semantics của từng phương thức.",
    "Tạo SupabaseRepository và FakeRepository cùng tuân thủ contract.",
    "Viết contract test dùng chung để bảo đảm implementation mới không làm thay đổi hành vi của service.",
]: bullet(doc, x)
heading(doc, "7.4.3.4. Interface Segregation Principle (ISP)", 2)
body(doc, "ISP yêu cầu caller chỉ phụ thuộc vào nhóm chức năng mà nó thật sự sử dụng. Dự án đã thể hiện nguyên lý này khá rõ ở customer API, component props và các domain router. Một số service lớn có thể tiếp tục được chia nhỏ khi quy mô tăng.")
table(doc, ["Nền tảng hiện tại", "Điểm tốt", "Hướng mở rộng"], [
    ["Customer service modules", "Appointment, boarding, medication, notification, pet và profile đã tách riêng.", "Giữ contract nhỏ và bổ sung schema dùng chung."],
    ["Doctor/staff feature service", "Đã gom theo đúng portal và che giấu HTTP khỏi phần lớn component.", "Tách tiếp ExamApi, BoardingApi, PaymentApi khi số use case tăng."],
    ["Backend domain modules", "Auth, customer, doctor, staff và admin có public router/service rõ.", "Tách admin report/settings/catalog thành service chuyên biệt."],
], [4.0, 6.0, 6.0], font_size=9.7)
body(doc, "Khi tách interface/service theo use case, component hoặc route chỉ import phần cần dùng. Điều này làm giảm phụ thuộc dư thừa, dễ mock khi test và giúp quyền truy cập nghiệp vụ rõ ràng hơn.")

# PAGE 12 - DIP AND IDEAL PROPERTIES
page_break(doc)
heading(doc, "7.4.3.5. Dependency Inversion Principle (DIP)", 2)
body(doc, "DIP yêu cầu logic cấp cao phụ thuộc abstraction thay vì chi tiết hạ tầng. Pet Service đã có service layer giữa UI và backend, đồng thời có lib cho Supabase/Prisma/JWT. Tuy nhiên đây mới là bước đầu vì business service vẫn import chi tiết hạ tầng trực tiếp.")
label(doc, "Phụ thuộc hiện tại:")
code_block(doc, "Route -> Service -> Supabase SDK -> PostgreSQL\n                    |\n                    +-> Email service -> Nodemailer\n                    +-> Storage -> Supabase Storage")
label(doc, "Thiết kế hướng DIP đề xuất:")
code_block(doc, "Use case service -> AppointmentRepository interface\n                 -> EmailSender interface\n                 -> FileStorage interface\n\nComposition root -> SupabaseAppointmentRepository\n                 -> NodemailerEmailSender\n                 -> SupabaseFileStorage")
body(doc, "Với cách thứ hai, service không cần biết dữ liệu đến từ Supabase hay Prisma, email được gửi qua Gmail hay provider khác. Khi unit test, nhóm truyền fake repository và fake email sender để kiểm tra nghiệp vụ mà không kết nối dịch vụ thật.")
heading(doc, "7.5. Đánh giá theo các tính chất mô-đun lý tưởng", 1)
table(doc, ["Tính chất", "Đánh giá", "Bằng chứng"], [
    ["Decomposable", "Khá", "Frontend/backend tách riêng; backend theo domain; frontend theo role/feature."],
    ["Composable", "Khá", "Page ghép component; app ghép router qua apiModules; service phối hợp theo luồng."],
    ["Understandable", "Khá", "Cấu trúc thư mục rõ nhưng file lớn làm tăng thời gian đọc."],
    ["Continuity", "Khá - đang hoàn thiện", "Phần lớn thay đổi nằm trong feature/domain; DTO và data-access có thể ổn định thêm."],
    ["Isolation", "Khá - đang hoàn thiện", "Middleware/lib/service đã cô lập nhiều chi tiết; event/repository sẽ tăng mức độc lập."],
], [3.0, 2.6, 10.4], font_size=9.8)
callout(doc, "Đánh giá chung", "Pet Service đáp ứng tốt phần lớn tiêu chí mô-đun: Decomposable, Composable và Understandable đều ở mức khá. Continuity và Isolation đã có nền tảng, cần được củng cố thêm bằng repository, schema/DTO ổn định và tách side effect.")

# PAGE 13 - RISKS
heading(doc, "7.6. Các điểm cần tiếp tục hoàn thiện", 1)
body(doc, "Qua quá trình tái cấu trúc, Pet Service đã hình thành nền tảng khá rõ về module, service layer, middleware và role-based frontend. Để giữ chất lượng này khi hệ thống tiếp tục mở rộng, nhóm xác định một số hạng mục nên hoàn thiện thêm. Các hạng mục dưới đây chiếm phần nhỏ hơn so với những phần thiết kế đã làm tốt.")
heading(doc, "7.6.1. Duy trì quy mô file hợp lý", 2)
body(doc, "Cấu trúc domain/feature đã giúp khoanh vùng tốt phần lớn chức năng. Bước tiếp theo là tách một số file còn dài theo use case để giảm merge conflict và giúp unit test tập trung hơn.")
heading(doc, "7.6.2. Tiếp tục làm mỏng Route", 2)
body(doc, "Nhiều route auth/customer đã gọi service khá rõ ràng. Nhóm có thể áp dụng đồng đều cách này cho medical, doctor, staff và services để route chỉ nhận request, gọi use case rồi trả response.")
heading(doc, "7.6.3. Củng cố ranh giới Data-access", 2)
body(doc, "Supabase client hiện đã được khởi tạo tập trung trong lib. Việc bổ sung repository theo domain sẽ giúp giữ tên bảng, quan hệ join và cấu trúc row ở một lớp thấp hơn, qua đó bảo vệ business service khi schema thay đổi.")
heading(doc, "7.6.4. Thống nhất Contract API", 2)
body(doc, "Customer API đã có requestJson và TypeScript type làm nền tảng tốt. Có thể mở rộng cùng chuẩn này sang doctor, staff và admin; response domain nên tách khỏi màu sắc hoặc Tailwind class của giao diện web.")
heading(doc, "7.6.5. Tăng khả năng cô lập Side effect", 2)
body(doc, "Email, PDF và storage đã có service riêng thay vì nằm hoàn toàn trong route. Bước hoàn thiện là đặt interface/event ở trước các service này để retry, test và thay provider thuận tiện hơn.")
table(doc, ["Hạng mục hoàn thiện", "Nền tảng đã có", "Bước tiếp theo"], [
    ["Tách file lớn", "Đã có ranh giới domain và feature.", "Tách tiếp theo use case trước khi thêm tính năng lớn."],
    ["Data-access", "Supabase client được tạo tập trung trong lib.", "Bổ sung repository ở domain trọng yếu."],
    ["API contract", "Customer API đã có requestJson và TypeScript DTO.", "Chuẩn hóa ApiResult/ApiError cho toàn frontend."],
    ["Presentation mapping", "Backend đã trả status/label nhất quán ở nhiều luồng.", "Chuyển màu/class cụ thể về frontend."],
], [6.0, 4.0, 6.0], font_size=9.3)

# PAGE 14 - ROADMAP
heading(doc, "7.7. Lộ trình cải thiện thiết kế", 1)
body(doc, "Lộ trình dưới đây không yêu cầu viết lại dự án. Mục tiêu là cải thiện dần theo rủi ro, giữ nguyên chức năng đang chạy và tạo ranh giới tốt hơn cho phần phát triển tiếp theo.")
heading(doc, "7.7.1. Giai đoạn ngắn hạn", 2)
for x in [
    "Thống nhất một ApiClient cho frontend: apiUrl, auth header, JSON, timeout, ApiError và invalid-session.",
    "Tách medical.routes.js thành route mỏng và service theo exam, prescription, specialist order, notification.",
    "Tách component trên 700-800 dòng thành container, hook, view và modal.",
    "Chuẩn hóa response theo { ok, data, message } hoặc contract tương đương.",
    "Chuyển Tailwind class và màu trạng thái khỏi backend về frontend.",
]: numbered(doc, x)
heading(doc, "7.7.2. Giai đoạn trung hạn", 2)
for x in [
    "Tách admin.service.js và staff.service.js theo domain/use case.",
    "Tạo AppointmentRepository, PetRepository, BoardingRepository và InvoiceRepository.",
    "Bổ sung validation schema tại API boundary cho input quan trọng.",
    "Tạo global error middleware và error type chung cho backend.",
    "Chuyển ownership thật cho reports, grooming và billing thay vì chỉ re-export.",
]: numbered(doc, x)
heading(doc, "7.7.3. Giai đoạn dài hạn", 2)
for x in [
    "Áp dụng dependency injection tại composition root cho repository, email, storage và clock.",
    "Dùng domain/application event cho AppointmentConfirmed, ExamCompleted, InvoicePaid và BoardingUpdated.",
    "Dùng OpenAPI hoặc shared schema để đồng bộ contract frontend-backend.",
    "Viết contract test cho repository và integration test cho các luồng nghiệp vụ chính.",
]: numbered(doc, x)
heading(doc, "7.7.4. Tiêu chí hoàn thành", 2)
table(doc, ["Tiêu chí", "Kết quả mong đợi"], [
    ["Route mỏng", "Không truy vấn database trực tiếp; chỉ validate, gọi use case và trả response."],
    ["Service tập trung", "Mỗi service có một nhóm use case liên quan và quy mô đọc được."],
    ["Data access thay thế được", "Business logic chạy với repository thật hoặc fake."],
    ["Frontend API thống nhất", "Mọi feature dùng cùng auth/error/JSON behavior."],
    ["Side effect cô lập", "Email/storage lỗi có cơ chế retry hoặc xử lý riêng."],
], [5.2, 10.8], font_size=10.1)

# PAGE 15 - CONCLUSION
page_break(doc)
heading(doc, "7.8. Kết luận", 1)
body(doc, "Dựa trên các tiêu chí của tài liệu Good Design, Pet Service đã có nền tảng thiết kế tương đối tốt. Việc tách frontend và backend, chia frontend theo role/feature, chia backend theo domain module, tách app khỏi server và đưa config, middleware, lib vào khu vực riêng giúp cấu trúc dự án dễ hiểu hơn so với cách tổ chức routes/services phẳng trước đây.")
body(doc, "Hệ thống thể hiện khá rõ nhiều nguyên lý quan trọng: SRP qua ranh giới app/server và các service nhỏ; OCP qua apiModules và khả năng ghép component; ISP qua customer API chuyên biệt; DIP ở mức nền tảng qua service layer, lib và middleware dùng chung. TypeScript DTO cùng public module contract cũng tạo tiền đề tốt để tiến tới LSP theo hướng implementation có thể thay thế.")
body(doc, "Nhóm 23 đánh giá khoảng 70% tiêu chí Good Design đã được triển khai tốt và có bằng chứng trực tiếp trong source. Khoảng 30% còn lại là công việc hoàn thiện, không phải yêu cầu viết lại hệ thống: tách một số file lớn, bổ sung repository, chuẩn hóa DTO/HTTP client và cô lập side effect. Khi các bước này được thực hiện, dự án sẽ dễ bảo trì, kiểm thử và mở rộng hơn nữa.")
callout(doc, "Kết quả tổng hợp", "Khoảng 70% tiêu chí Good Design đã được thể hiện tốt trong code hiện tại: modularity rõ, service/middleware dùng chung, frontend theo role-feature và backend theo domain. Khoảng 30% còn lại là các bước hoàn thiện về file lớn, repository, contract API và side effect.")
heading(doc, "Tài liệu và mã nguồn đối chiếu", 2)
for x in [
    "Tài liệu Thiết kế tốt.pdf - tài liệu tham khảo về Modularity, Coupling, Cohesion và SOLID.",
    "README.md, docs/architecture.mdx, docs/frontend.mdx, docs/backend.mdx.",
    "backend/src/app.js, server.js, config/, middlewares/, lib/, modules/ và utils/.",
    "frontend/src/app, components, features, services, types và utils.",
]: bullet(doc, x)
heading(doc, "Thành viên thực hiện", 2)
table(doc, ["STT", "Họ và tên", "Nhóm"], [
    ["1", "Trần Đức Nam Anh", "23"],
    ["2", "Nguyễn Đức Hiếu", "23"],
    ["3", "Triệu Trường Giang", "23"],
], [2.0, 10.0, 4.0], font_size=10.8)

doc.core_properties.title = "Chương 7 - Tài liệu Good Design - Pet Service - Nhóm 23"
doc.core_properties.subject = "Modularity, Coupling, Cohesion và SOLID"
doc.core_properties.author = "Nhóm 23 - Trần Đức Nam Anh, Nguyễn Đức Hiếu, Triệu Trường Giang"
doc.core_properties.keywords = "Pet Service, Good Design, Modularity, Coupling, Cohesion, SOLID, ITSS"
doc.save(OUTPUT)
print(OUTPUT)
